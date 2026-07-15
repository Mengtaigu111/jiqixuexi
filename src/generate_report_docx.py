"""Generate a Word (.docx) version of the report from the SAME honest content
source as the PDF.

This deliberately does NOT reuse the legacy ``scripts/generate_word_report.py``,
which hard-codes the old (retracted) "DMSAFormer wins at 365 days" narrative.
Instead it pulls text from ``build_report_text_pages`` (the raw-口径, honest
version shared with the PDF), the raw ``summary.csv`` main table, the metric/
prediction comparison figures, and the four diagnostic figures. So the DOCX and
the PDF are guaranteed to tell the same story.
"""

from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.generate_report_pdf import build_report_text_pages, _display_model

ROOT = Path(__file__).resolve().parents[1]
FONT = "Microsoft YaHei"
DARK = RGBColor(0x17, 0x36, 0x5D)
REPO_URL = "https://github.com/Mengtaigu111/jiqixuexi"


def _set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    for name, size in [("Heading 1", 17), ("Heading 2", 13)]:
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = DARK


def _add_body(doc: Document, text: str) -> None:
    """Render a page body: split code blocks (indented pseudo-code) from prose."""
    for block in text.strip().split("\n\n"):
        block = block.strip()
        if not block:
            continue
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.2
        run = para.add_run(block)
        run.font.name = FONT
        run.font.size = Pt(10.5)


def _is_heading(title: str) -> int:
    """Return heading level for a page title (1 for top sections, 2 for sub)."""
    head = title.strip()
    if head[:2] in {"1.", "2.", "3.", "4."} and (len(head) < 4 or head[1] == "."):
        # e.g. "1. 问题介绍", "4. 讨论与后续工作"
        if len(head) > 2 and head[2] == " ":
            return 1
    if head[:3] in {"1.1", "1.2", "2.1", "2.2", "3.1", "3.2", "3.3", "3.4", "3.5", "3.6"}:
        return 2
    return 1


def _summary_table(doc: Document, summary: pd.DataFrame, horizon: int) -> None:
    subset = summary[summary["Horizon"].astype(int) == horizon].sort_values("MSE mean")
    doc.add_paragraph().add_run(f"表  {horizon} 天预测结果（原尺度，未校准）").bold = True
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ["模型", "MSE mean ± std", "MAE mean ± std", "Runs"]):
        cell.text = label
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
    for _, row in subset.iterrows():
        cells = table.add_row().cells
        cells[0].text = _display_model(str(row["Model"]))
        cells[1].text = f"{float(row['MSE mean']):.2f} ± {float(row['MSE std']):.2f}"
        cells[2].text = f"{float(row['MAE mean']):.2f} ± {float(row['MAE std']):.2f}"
        cells[3].text = str(int(row["Runs"]))
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)


def _add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4F, 0x62, 0x6F)


def generate_docx(report_dir: str | Path = "report", results_dir: str | Path = "results") -> Path:
    report_path = Path(report_dir)
    results_path = Path(results_dir)
    report_path.mkdir(parents=True, exist_ok=True)
    summary_path = results_path / "metrics" / "summary.csv"
    figures_dir = results_path / "figures"

    pages = build_report_text_pages(summary_path)
    summary = pd.read_csv(summary_path)

    doc = Document()
    _set_base_style(doc)

    # Cover
    cover = pages[0]
    title = doc.add_heading(cover.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("LSTM、Transformer 与 DMSAFormer 的公平三模型比较研究")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x4F, 0x62, 0x6F)
    _add_body(doc, cover.body)
    doc.add_page_break()

    # Body sections. Insert the raw summary table right after "3. 结果与分析"
    # and the diagnostic figures right after "3.5" diagnostic-two text page.
    for page in pages[1:]:
        level = _is_heading(page.title)
        doc.add_heading(page.title, level=level)
        _add_body(doc, page.body)

        if page.title.startswith("3. "):
            _summary_table(doc, summary, 90)
            _summary_table(doc, summary, 365)
            _add_figure(doc, figures_dir / "metric_bar_mse.png", "图  MSE 指标对比（90/365 天）")
            _add_figure(doc, figures_dir / "metric_bar_mae.png", "图  MAE 指标对比（90/365 天）")
            _add_figure(doc, figures_dir / "prediction_comparison_90.png", "图  90 天预测曲线对比")
            _add_figure(doc, figures_dir / "prediction_comparison_365.png", "图  365 天预测曲线对比")

        if page.title.startswith("3.5"):
            _add_figure(doc, figures_dir / "diag_naive_floor.png", "图  深度模型 vs 朴素基线地板")
            _add_figure(doc, figures_dir / "diag_capacity_vs_error.png", "图  参数量 vs 测试误差")
            _add_figure(doc, figures_dir / "diag_overfitting_gap.png", "图  过拟合诊断（验证/训练损失比）")
            _add_figure(doc, figures_dir / "diag_ablation.png", "图  DMSAFormer 逐模块消融")

    out_path = report_path / "ML_household_power_report.docx"
    doc.save(str(out_path))
    return out_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the Word (.docx) report from the honest raw content source.")
    parser.add_argument("--report_dir", default="report")
    parser.add_argument("--results_dir", default="results")
    return parser.parse_args()


def main() -> None:
    out = generate_docx(**vars(parse_args()))
    print(f"Saved DOCX report: {out}")


if __name__ == "__main__":
    main()
